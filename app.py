

######################### REAEDME.md #####################################

### license : Creative Commons Attribution-NonCommercial-NoDerivatives (CC BY-NC-ND)
### https://creativecommons.org/licenses/by-nc-nd/4.0/
### Made by Woo Yoon Kyu


#########################################################################

from docx import Document
from docx.shared import Pt, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.text import WD_ALIGN_PARAGRAPH
import streamlit as st
from langgraph.graph import StateGraph
from typing import TypedDict, Literal
import fitz  # PyMuPDF
import tempfile
from graph import *  # 요약 & 평가 로직
from dotenv import load_dotenv
import os
import re
import copy

load_dotenv()






# ====== STREAMLIT 설정 ======
st.set_page_config(page_title="논문 요약 멀티 에이전트", layout="wide")

# ✅ 스타일: 최대 너비 제한 (CSS 삽입)
# Streamlit 상단에 CSS 추가
st.markdown("""
<style>
    .main {
        max-width: 1000px;
        margin: 0 auto;
        padding-top: 2rem;
    }
    .element-container textarea, .element-container pre, .element-container code {
        white-space: pre-wrap !important;
        word-wrap: break-word !important;
        overflow-x: auto !important;
    }
    .stTextArea > div > textarea {
        font-family: '맑은 고딕', sans-serif;
        font-size: 0.92rem;
        line-height: 1.4;
    }
</style>
""", unsafe_allow_html=True)



#################


BASE_DIR = os.path.dirname(os.path.abspath(__file__))  # 현재 파일 기준 디렉토리
logo_path = os.path.join(BASE_DIR, "IIPL.PNG")
license_path = os.path.join(BASE_DIR, "license.png")

with st.container():
    col1, col2 = st.columns([6, 1])
    with col1:
        st.title("🧠 논문 요약 멀티 에이전트")
        st.image(logo_path, width=1000)  # 로고 경로 수정 필요
    with col2:
        st.image(license_path, width=200)

st.markdown("""
#### 🤖 Summarizer & Critic 협업 기반 논문 요약기
- 논문을 **최대 5회 요약-평가 루프**를 통해 개선합니다.
- 일반 : **기법 설명 / 장점 및 효과 / 예상 질문 & 답변** 을 추출합니다.
- 전문가 : arxiv 논문 스타일로 논문의 전문가적 설명을 제시합니다.
- **과제 목적**에 따라서 해당 논문을 분석, 일반 문서를 생성해 냅니다.
- 반대로 논문으로부터 과제 목적 자체 추출이 가능하고, 그를 기반으로 다시 분석이 가능합니다.
- **과제의 목적, 진행 상황(예시 참조) 을 입력 시** 그에 맞게 해당 논문의 방법론이 어떻게 적용될 수 있을지 추천하여 줍니다.
- 목적 및 진행 상황을 입력하지 않으면 단순히 논문 요약 Task만 수행합니다.
""")

# ======= 자동 목적 제안 =======
def suggest_project_goal_from_summaries(summary: str, expanded: str) -> str:
    prompt = f"""
다음은 논문 요약 및 쉬운 설명이야. 이 내용을 바탕으로 이 기술이 어떤 과제나 프로젝트 목적에 활용될 수 있을지 제안해줘.

- 문장은 1~2문장으로 간단히.
- 산업 현장, 공공 서비스, 사용자 편의성, 자동화, 비용 절감 등 실용적인 문제 해결 관점에서 써줘.

===== 전문가 요약 =====
{summary}

===== 일반인용 설명 =====
{expanded}
"""
    return call_openai_gpt(prompt)



# ======= md, WORD 저장 =======
def save_summary_to_md(text: str, title: str) -> str:
    md_content = f"# {title}\n\n" + text.strip()
    temp_file = tempfile.NamedTemporaryFile(delete=False, suffix=".md", mode="w", encoding="utf-8")
    temp_file.write(md_content)
    temp_file.close()
    return temp_file.name


def save_summary_to_word(summary: str, title: str) -> str:
    doc = Document()

    # 제목
    title_para = doc.add_paragraph()
    run = title_para.add_run(title)
    run.font.size = Pt(16)
    run.bold = True
    run.font.color.rgb = RGBColor(0xC0, 0x00, 0x00)
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_para.add_run("\n" + "―" * 50)

    # 본문 스타일 설정
    style = doc.styles['Normal']
    font = style.font
    font.name = '맑은 고딕'
    font.size = Pt(11)

    # ✅ 한글 깨짐 방지 설정
    rFonts = style.element.rPr.rFonts
    rFonts.set(qn('w:eastAsia'), '맑은 고딕')

    for line in summary.strip().split("\n"):
        if line.strip() == "":
            doc.add_paragraph("")
        elif line.startswith("1)") or line.startswith("2)") or line.startswith("3)") or line.startswith("4)"):
            para = doc.add_paragraph()
            run = para.add_run(line)
            run.bold = True
            run.font.color.rgb = RGBColor(0x2E, 0x74, 0xB5)
        elif line.startswith("(1)") or line.startswith("(2)") or line.startswith("(3)") or line.startswith("(4)"):
            para = doc.add_paragraph()
            run = para.add_run(line)
            run.bold = True
            run.font.color.rgb = RGBColor(0x2E, 0x74, 0xB5)
        else:
            doc.add_paragraph(line)

    temp_file = tempfile.NamedTemporaryFile(delete=False, suffix=".docx")
    doc.save(temp_file.name)
    return temp_file.name

# ======= 파일 업로드 =======
if "state" not in st.session_state:
    st.session_state.state = None
    st.session_state.title = ""
    st.session_state.body = ""

uploaded_file = st.file_uploader("📤 논문 PDF 업로드", type=["pdf"])


if uploaded_file:
    # 파일 이름이 바뀌었는지 확인
    new_file_name = uploaded_file.name
    if st.session_state.get("last_uploaded_file") != new_file_name:
        title, body = extract_text_from_pdf(uploaded_file)
        st.session_state.title = title
        st.session_state.body = body
        st.session_state.state = None
        st.session_state.last_uploaded_file = new_file_name
        # st.session_state.project_goal = ""
        # st.session_state.project_status = ""
        # st.session_state.suggested_goal = ""

    with st.spinner("📚 텍스트 추출 중..."):
        st.success("✅ 텍스트 추출 완료!")

    with st.expander("📄 논문 미리보기 (일부)", expanded=False):
        st.text_area("본문", st.session_state.body[:3000], height=200)

    # ✅ 과제 목적 입력창 추가
    project_goal = st.text_area("🎯 이 논문이 활용될 과제의 목적을 입력하세요", height=150)
    st.session_state.project_goal = project_goal  # 상태에 저장

    project_status = st.text_area(
        "📅 과제의 진행 상황을 입력해주세요",
        help="아래 형식을 따르세요:\n- [완료된 단계 (~MM/DD)] : 이미 수행된 작업\n- [예정된 단계 (~MM/DD)] : 앞으로 수행할 계획인 작업",
        placeholder="""
    📌 입력 형식 예시:
    - [완료된 단계 (~MM/DD)] : 이미 수행한 작업
    - [예정된 단계 (~MM/DD)] : 앞으로 수행할 계획인 작업

    [데이터 수집 및 실험 완료 (~4/6)]  # 완료된 단계
    <QA 시스템>
    - Dataset: snuh/ClinicalQA
    - 사용 모델: Exaone-7.8B-instruct, EEVE-10.8B, QWEN 2.5 7B, OpenBioLLM + chatvector + Blocking + LoRA

    <Dialogue Generation>
    - Dataset: 내부 커스텀 데이터셋
    - 사용 모델: GPT-4o

    [의료 도메인 데이터 추가 크롤링 (~4/13)]  # 완료된 단계
    - 대한고혈압학회, 당뇨병학회, 치매 관련 협회 등에서 정보 수집
    - 기존 QA 구조에 통합 가능하도록 전처리

    [ko-LLM 대상 기법 적용 및 실험 (~4/27)]  # 예정된 단계
    - 예정 기법:
    - Chat Vector
    - 특수문자 Unicode 블로킹
    - CoT Steering
    - Self-Consistency (또는 기타 파인튜닝 기법)
    """,
        height=700
    )

    st.session_state.project_status = project_status


if st.session_state.body and st.button("🚀 요약 시작"):
    with st.spinner("요약-비평 루프 진행 중 (최대 5회)..."):
        graph = build_graph()

        state = {
            "document": st.session_state.body,
            "summary": "",
            "expanded_summary": "",
            "feedback": "",
            "critic_result": "",
            "status": "in_progress",
            "loop_count": 0,
            "project_goal": st.session_state.get("project_goal", "").strip(),
            "project_status": st.session_state.get("project_status", "").strip()
        }
                
        loop_container = st.container()
        loop_placeholder = loop_container.empty()

        while state["status"] != "done" and state["loop_count"] < 5:
            with loop_placeholder.container():
                st.markdown(f"### 🔁 {state['loop_count']+1}회차 요약/평가")
                st.markdown("#### Critic 상태")
                st.markdown(state["critic_result"][:300] + "...")
            state = graph.invoke(copy.deepcopy(state))


        st.session_state.state = state
        st.success("🎉 최종 요약 완료!")
######################




# ===== 결과 표시 & 다운로드 =====
if st.session_state.state:
    state = st.session_state.state


    if "alignment_analysis" in state:
        st.subheader("🔗 과제 정렬 분석 결과 (논문 방법론과 과제 진행의 연결성)")
        st.code(state["alignment_analysis"], language="markdown")

        # Word 다운로드
        alignment_word = save_summary_to_word(state["alignment_analysis"], f"{st.session_state.title} - 과제 정렬 분석")
        with open(alignment_word, "rb") as f:
            st.download_button("📥 과제 정렬 분석 Word 다운로드", f, file_name="과제_정렬_분석.docx")

        # Markdown 다운로드
        alignment_md = save_summary_to_md(state["alignment_analysis"], f"{st.session_state.title} - 과제 정렬 분석")
        with open(alignment_md, "rb") as f:
            st.download_button("📥 과제 정렬 분석 Markdown 다운로드", f, file_name="과제_정렬_분석.md")

    if "summary_pro_3line" in state or "summary_nonpro_3line" in state.keys():
        st.markdown("## 🧾 최종 3줄 요약")

        if "summary_pro_3line" in state:
            st.markdown("#### 📌 전문가용 요약 ")
            st.markdown(state["summary_pro_3line"])

            md_file_pro = save_summary_to_md(state["summary_pro_3line"], f"{st.session_state.title} - 전문가용 요약")
            with open(md_file_pro, "rb") as f:
                st.download_button("⬇️ 전문가용 요약 다운로드 (.md)", f, file_name="전문가용_3줄_요약.md")

        if "summary_nonpro_3line" in state:
            st.markdown("#### 💡 일반인용 요약 (3줄)")
            st.markdown(state["summary_nonpro_3line"])

            md_file_nonpro = save_summary_to_md(state["summary_nonpro_3line"], f"{st.session_state.title} - 일반인용 요약")
            with open(md_file_nonpro, "rb") as f:
                st.download_button("⬇️ 일반인용 요약 다운로드 (.md)", f, file_name="일반인용_3줄_요약.md")

    # 기존 내용 계속

    st.subheader("📌 1차 요약 결과 (전문가용)")
    st.code(state["summary"], language="markdown")

    st.subheader("💬 2차 쉬운 설명 (일반인용)")
    st.code(state["expanded_summary"], language="markdown")


    if state.get("critic_result"):
        st.subheader("🧐 Critic 최종 평가 결과")
        st.text_area("Critic 평가", state["critic_result"], height=700)

    word_summary = save_summary_to_word(state["summary"], f"{st.session_state.title} - 1차 요약")
    word_expanded = save_summary_to_word(state["expanded_summary"], f"{st.session_state.title} - 2차 쉬운 설명")

    st.markdown("### 📥 결과 다운로드")
    col1, col2 = st.columns(2)

    with col1:
        st.markdown("#### 🧠 1차 요약")
        with open(word_summary, "rb") as f:
            st.download_button("📝 Word 다운로드", f, file_name="1차_요약.docx")

    with col2:
        st.markdown("#### 💬 2차 설명")
        with open(word_expanded, "rb") as f:
            st.download_button("📝 Word 다운로드", f, file_name="2차_설명.docx")

    if state["loop_count"] >= 5 and state["status"] != "done":
        with open("critic_result.txt", "w", encoding="utf-8") as f:
            f.write(state["critic_result"])
        with open("critic_result.txt", "rb") as f:
            st.download_button("🧐 Critic 평가 결과 다운로드", f, file_name="critic_result.txt")








    # 제안 다시 받기 버튼
    if st.button("🔄 제안 다시 받기"):
        with st.spinner("GPT가 새로운 과제 목적을 제안 중입니다..."):
            new_goal = suggest_project_goal_from_summaries(
                state["summary"], state["expanded_summary"]
            )
            st.session_state.suggested_goal = new_goal
            st.session_state.project_goal = new_goal
        st.success("✅ 새로운 과제 목적이 제안되었습니다!")

    # 보여주기
    st.text_area("💡 GPT가 제안한 과제 목적", st.session_state.get("suggested_goal", ""), height=100)

    # 재요약 실행
    if st.button("🚀 이 제안 목적으로 재요약 실행"):
        with st.spinner("새로운 목적 기준으로 재요약 중..."):
            graph = build_graph()

            new_state = {
            "document": st.session_state.body,
            "summary": "",
            "expanded_summary": "",
            "feedback": "",
            "critic_result": "",
            "status": "in_progress",
            "loop_count": 0,
            "project_goal": st.session_state.get("project_goal", "").strip(),
            "project_status": st.session_state.get("project_status", "").strip()
            }
            

            new_loop = st.container()
            new_loop_placeholder = new_loop.empty()
            while new_state["status"] != "done" and new_state["loop_count"] < 5:
                with new_loop_placeholder.container():
                    st.markdown(f"### 🔁 재요약 {new_state['loop_count']+1}회차")
                    st.markdown("#### Critic 상태")
                    st.markdown(new_state["critic_result"][:300] + "...")

                new_state = graph.invoke(copy.deepcopy(new_state))

            st.session_state.state = new_state
            st.success("✅ 제안 목적 기준으로 재요약 완료!")

            print(new_state)



